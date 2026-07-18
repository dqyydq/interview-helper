import io
import shutil
import uuid
from pathlib import Path

import pytest
import pytest_asyncio
from docx import Document
from httpx import ASGITransport, AsyncClient
from pypdf import PdfWriter
from sqlalchemy import delete

from app.core.config import settings
from app.db.models.job import BackgroundJob
from app.db.models.resume import Resume
from app.db.session import async_session_factory, engine
from app.main import app
from app.workers.resume_jobs import run_once


async def clear_resume_data() -> None:
    async with async_session_factory() as session:
        await session.execute(delete(BackgroundJob))
        await session.execute(delete(Resume))
        await session.commit()


def make_test_root() -> Path:
    backend_root = Path(__file__).resolve().parents[2]
    test_root = backend_root / ".test-data" / uuid.uuid4().hex
    test_root.mkdir(parents=True)
    return test_root


@pytest_asyncio.fixture(autouse=True)
async def isolated_resume_data(monkeypatch: pytest.MonkeyPatch):
    test_root = make_test_root()
    monkeypatch.setattr(settings, "upload_dir", test_root / "uploads")
    monkeypatch.setattr(settings, "resume_upload_max_bytes", 5 * 1024 * 1024)
    monkeypatch.setattr(settings, "job_poll_interval_seconds", 0.01)
    await clear_resume_data()
    yield
    await clear_resume_data()
    await engine.dispose()
    shutil.rmtree(test_root)


def docx_bytes() -> bytes:
    document = Document()
    document.add_heading("项目经历", level=1)
    document.add_paragraph("构建支持上下文压缩的多 Agent 面试系统")
    document.add_heading("专业技能", level=1)
    document.add_paragraph("FastAPI、PostgreSQL、RAG")
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def blank_pdf_bytes() -> bytes:
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    stream = io.BytesIO()
    writer.write(stream)
    return stream.getvalue()


@pytest.mark.asyncio
async def test_text_resume_upload_is_idempotent_and_worker_parses_sections_and_claims() -> None:
    content = """个人简介
大模型应用开发工程师

项目经历
- 构建支持长对话上下文压缩的模拟面试系统
- 使用 FastAPI 与 PostgreSQL 实现任务队列

专业技能
Python、RAG、Agent、SSE
""".encode()
    async with AsyncClient(
        transport=ASGITransport(app=app, raise_app_exceptions=False),
        base_url="http://test",
    ) as client:
        uploaded = await client.post(
            "/api/resumes",
            files={"file": ("resume.txt", content, "text/plain")},
        )
        duplicate = await client.post(
            "/api/resumes",
            files={"file": ("resume-again.txt", content, "text/plain")},
        )

        assert uploaded.status_code == 201
        assert uploaded.json()["reused"] is False
        assert uploaded.json()["job"]["status"] == "queued"
        assert duplicate.status_code == 201
        assert duplicate.json()["reused"] is True
        assert duplicate.json()["resume"]["id"] == uploaded.json()["resume"]["id"]
        assert duplicate.json()["job"]["id"] == uploaded.json()["job"]["id"]

        assert await run_once("test-worker") is True
        resume_id = uploaded.json()["resume"]["id"]
        job_id = uploaded.json()["job"]["id"]
        detail = await client.get(f"/api/resumes/{resume_id}")
        job = await client.get(f"/api/jobs/{job_id}")
        events = await client.get(f"/api/jobs/{job_id}/events")

    assert detail.status_code == 200
    assert detail.json()["parse_status"] == "ready"
    assert [section["section_type"] for section in detail.json()["sections"]] == [
        "summary",
        "projects",
        "skills",
    ]
    assert len(detail.json()["claims"]) >= 4
    assert "storage_path" not in detail.text
    assert job.json()["status"] == "completed"
    assert job.json()["progress"] == 1.0
    assert "event: job" in events.text
    assert f"id: {job.json()['version']}" in events.text


@pytest.mark.asyncio
async def test_docx_and_markdown_resumes_parse_without_model_calls() -> None:
    async with AsyncClient(
        transport=ASGITransport(app=app, raise_app_exceptions=False),
        base_url="http://test",
    ) as client:
        docx_upload = await client.post(
            "/api/resumes",
            files={
                "file": (
                    "resume.docx",
                    docx_bytes(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        markdown_upload = await client.post(
            "/api/resumes",
            files={"file": ("resume.md", b"# Skills\nPython and RAG", "text/markdown")},
        )
        assert docx_upload.status_code == 201
        assert markdown_upload.status_code == 201
        assert await run_once("test-worker") is True
        assert await run_once("test-worker") is True
        docx_detail = await client.get(
            f"/api/resumes/{docx_upload.json()['resume']['id']}"
        )
        markdown_detail = await client.get(
            f"/api/resumes/{markdown_upload.json()['resume']['id']}"
        )

    assert docx_detail.json()["parse_status"] == "ready"
    assert markdown_detail.json()["parse_status"] == "ready"
    assert any(item["section_type"] == "projects" for item in docx_detail.json()["sections"])
    assert any(item["section_type"] == "skills" for item in markdown_detail.json()["sections"])


@pytest.mark.asyncio
async def test_empty_pdf_reaches_safe_terminal_failure_and_sse_resumes_from_event_id() -> None:
    async with AsyncClient(
        transport=ASGITransport(app=app, raise_app_exceptions=False),
        base_url="http://test",
    ) as client:
        uploaded = await client.post(
            "/api/resumes",
            files={"file": ("blank.pdf", blank_pdf_bytes(), "application/pdf")},
        )
        assert uploaded.status_code == 201
        assert await run_once("test-worker") is True
        job_id = uploaded.json()["job"]["id"]
        job = await client.get(f"/api/jobs/{job_id}")
        resume = await client.get(f"/api/resumes/{uploaded.json()['resume']['id']}")
        replay = await client.get(
            f"/api/jobs/{job_id}/events",
            headers={"Last-Event-ID": str(job.json()["version"])},
        )

    assert job.json()["status"] == "failed"
    assert job.json()["error_code"] == "resume_text_empty"
    assert resume.json()["parse_status"] == "failed"
    assert replay.text == ""


@pytest.mark.asyncio
async def test_resume_upload_rejects_unsupported_mime_signature_size_and_empty_text(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    async with AsyncClient(
        transport=ASGITransport(app=app, raise_app_exceptions=False),
        base_url="http://test",
    ) as client:
        unsupported = await client.post(
            "/api/resumes",
            files={"file": ("resume.exe", b"binary", "application/octet-stream")},
        )
        mismatch = await client.post(
            "/api/resumes",
            files={"file": ("resume.pdf", b"%PDF-fake", "text/plain")},
        )
        bad_signature = await client.post(
            "/api/resumes",
            files={"file": ("resume.pdf", b"not-a-pdf", "application/pdf")},
        )
        empty = await client.post(
            "/api/resumes",
            files={"file": ("resume.txt", b"  \n", "text/plain")},
        )
        monkeypatch.setattr(settings, "resume_upload_max_bytes", 8)
        too_large = await client.post(
            "/api/resumes",
            files={"file": ("resume.txt", b"more than eight bytes", "text/plain")},
        )

    assert unsupported.status_code == 415
    assert unsupported.json()["code"] == "resume_type_unsupported"
    assert mismatch.status_code == 415
    assert mismatch.json()["code"] == "resume_mime_mismatch"
    assert bad_signature.status_code == 422
    assert bad_signature.json()["code"] == "resume_signature_invalid"
    assert empty.status_code == 422
    assert empty.json()["code"] == "resume_file_empty"
    assert too_large.status_code == 413
    assert too_large.json()["code"] == "resume_file_too_large"
